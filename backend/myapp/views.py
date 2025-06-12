from django.http import HttpResponse
from django.contrib.auth import authenticate
from rest_framework import generics, views, status
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

from .models import CustomUser, TrainingPlan
from .serializers import UserSerializer, TrainingPlanSerializer

# Klasa do tworzenia użytkowników
class CreateUserView(generics.CreateAPIView):
    queryset = CustomUser.objects.all()
    serializer_class = UserSerializer

    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        if CustomUser.objects.filter(username=username).exists():
            return Response({'error': 'Użytkownik o tej nazwie już istnieje'}, status=status.HTTP_400_BAD_REQUEST)
        return super().post(request, *args, **kwargs)

# Widok logowania
class LoginView(views.APIView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        if user:
            return Response(UserSerializer(user).data)
        return Response({'error': 'Niepoprawne dane logowania'}, status=status.HTTP_400_BAD_REQUEST)

# Widok do dodawania planów treningowych
class TrainingPlanUploadView(views.APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        serializer = TrainingPlanSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Widok do pobierania treści planów treningowych
class TrainingPlanView(views.APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, plan_type):
        try:
            plan = TrainingPlan.objects.get(plan_type=plan_type)
            doc = Document(plan.document)
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            return Response({'plan': "\n".join(full_text)})
        except TrainingPlan.DoesNotExist:
            return Response({'error': 'Plan not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# Widok do generowania PDF z planem treningowym
class TrainingPlanPDFView(views.APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, plan_type):
        try:
            plan = TrainingPlan.objects.get(plan_type=plan_type)
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{plan_type}_training_plan.pdf"'
            
            p = canvas.Canvas(response, pagesize=letter)
            doc = Document(plan.document)
            text = p.beginText(40, 750)
            text.setFont('Times-Roman', 12)
            for paragraph in doc.paragraphs:
                text.textLine(paragraph.text)
            p.drawText(text)
            p.showPage()
            p.save()
            return response
        except TrainingPlan.DoesNotExist:
            return Response({'error': 'Plan not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
